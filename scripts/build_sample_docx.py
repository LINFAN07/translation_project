"""產生範例 Word 譯文檔（供驗證用）。"""
from pathlib import Path

from docx import Document
from docx.shared import Pt


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "output"
    out_dir.mkdir(exist_ok=True)
    # 中文檔名 + 英文檔名（避免部分環境顯示亂碼）
    path_zh = out_dir / "茂爺爺的教誨_序言與第一章譯文.docx"
    path_en = out_dir / "translated_shigeru_prologue_ch1.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(11)

    doc.add_heading("《89歲現役交易員 大富豪茂爺爺的教誨》譯文（序言～第一章節選）", 0)
    doc.add_paragraph("來源：使用者提供之 PDF 節選；譯文為繁體中文，供專案品質驗證用。")

    doc.add_heading("關鍵術語對照", level=1)
    terms = [
        ("現役トレーダー", "現役交易員"),
        ("ネット取引", "網路交易"),
        ("板（いた）", "盤口／報價跳動表"),
        ("前場", "早盤"),
        ("寄り付き", "開盤"),
        ("約定", "成交"),
        ("終値", "收盤價"),
        ("東証プライム", "東證 Prime"),
        ("日経平均株価", "日經平均指數"),
        ("損切り", "停損"),
        ("配当金", "股利／分紅"),
    ]
    for ja, zh in terms:
        doc.add_paragraph(f"{ja} → {zh}", style="List Bullet")

    doc.add_heading("序言：沒想到這位老爺爺會改變我的命運（節選）", level=1)
    body1 = """「哈……今天又搞砸了嗎？」
我——伊藤慎平，一邊輕聲嘆氣，一邊緩緩邁開腳步。
從神戶的私立大學畢業後，進入這家大型二手車銷售公司已經 18 年了。同期的人有些 20 多歲就當上店長，現在甚至成了統括區域的區經理；反觀我的業績總是墊底，不知不覺中，後輩們一個個超越了我。

——我到底在做什麼呢……
我對汽車的熱情並未消退，知識也不輸給任何人。但就是賣不出去，拿不出成果。時間虛度，別說晉升店長，我連身為業務員的自信都快磨滅了。"""
    for p in body1.split("\n\n"):
        doc.add_paragraph(p.strip())

    doc.add_heading("第一章：僅僅一擊，就賺到一個月的薪水（節選）", level=1)
    body2 = """「看好了，這就是網路交易。」
上午 8 點。不知為何，我竟然待在一個初次見面的老爺爺家裡。
「這本手冊啊，是我的股票買賣紀錄。……小伙子，你對股票有興趣嗎？」

我誠實地回答：「並不是沒興趣，只是……有點害怕。萬一錢不見了怎麼辦？」
「那是因為做法太差勁了。股票可是很有趣的喔。」

上午 9 點。那一瞬間，正如茂爺爺所說，螢幕上顯示的數字一齊跳動了起來。光芒閃爍，數字變換，世界彷彿開始奔跑。接著——
「叮咚！」
一聲清脆的通知音在房間內響起。
「噢……剛才的委託成交了啊。」

茂爺爺立刻看向另一個螢幕。
「好耶，住友商事，全部成交！」
就在剛才，他委託的股票在市場上撮合成功。僅僅幾分鐘，他竟然賺到了 34 萬日圓的利潤。

那幾乎是我一個月實領的薪水。而他，僅僅靠著「一擊」就辦到了。
我呆然地盯著螢幕，感受到心臟劇烈地跳動著。我從來不知道，這世界上竟然有這樣的地方——"""
    for p in body2.split("\n\n"):
        doc.add_paragraph(p.strip())

    doc.add_paragraph("")
    doc.add_paragraph(
        "（說明：完整 PDF 篇幅較長，此檔為先前對話中已譯出之節選；"
        "若要整本自動輸出為 .docx，可再擴充 main.py 的輸出格式。）"
    )

    doc.save(path_zh)
    doc.save(path_en)
    print(f"已儲存：{path_zh}")
    print(f"已儲存：{path_en}")


if __name__ == "__main__":
    main()
