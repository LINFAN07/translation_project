"""
將翻譯結果寫入 Word (.docx)。
"""
import io
import os
from docx import Document
from docx.shared import Pt


def save_translation_docx(
    output_path: str,
    body: str,
    *,
    title: str | None = None,
    subtitle: str | None = None,
) -> None:
    """body 以雙換行分段寫入多個段落。"""
    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Microsoft JhengHei"
        doc.styles["Normal"].font.size = Pt(11)
    except Exception:
        pass

    if title:
        doc.add_heading(title, 0)
    if subtitle:
        doc.add_paragraph(subtitle)

    blocks = [b.strip() for b in body.split("\n\n") if b.strip()]
    for block in blocks:
        doc.add_paragraph(block)

    out_dir = os.path.dirname(os.path.abspath(output_path))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)


def build_translation_docx_bytes(
    body: str,
    *,
    title: str | None = None,
    subtitle: str | None = None,
) -> bytes:
    """在記憶體中建立 .docx，供 Streamlit 下載按鈕使用。"""
    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Microsoft JhengHei"
        doc.styles["Normal"].font.size = Pt(11)
    except Exception:
        pass

    if title:
        doc.add_heading(title, 0)
    if subtitle:
        doc.add_paragraph(subtitle)

    blocks = [b.strip() for b in body.split("\n\n") if b.strip()]
    for block in blocks:
        doc.add_paragraph(block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
